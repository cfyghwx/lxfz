import os

import docx
import jieba
import pandas as pd

#工具类，包含各种处理txt，csv和doc的方法
class DealtextUtil(object):
    # def __init__(self):

    def deal_text(self, text):
        a = text.find('如不服本判决')
        text = text[:a]
        return text

    #处理文本，删去标点符号等，并进行分词
    def readdoc(self,doc1):

        textc = ''
        for i in range(len(doc1.paragraphs)):
            text1 = doc1.paragraphs[i].text
            textc += text1
        textc = self.deal_text(textc)
        textwb = jieba.cut(textc, cut_all=False)
        str_out = ' '.join(textwb).replace('，', '').replace('。', '').replace('？', '').replace('！', '') \
            .replace('“', '').replace('”', '').replace('：', '').replace('…', '').replace('（', '').replace('）', '') \
            .replace('—', '').replace('《', '').replace('》', '').replace('、', '').replace('‘', '') \
            .replace('’', '')
        list = str_out.split(' ')
        while '' in list:
            list.remove('')
        return list

    # 不分词返回文本结果
    def readdoc2(self,path):
        doc1 = docx.Document(path)
        # print('1',type(doc1))
        textc = ''
        for i in range(len(doc1.paragraphs)):
            text1 = doc1.paragraphs[i].text
            textc += text1
        textc = self.deal_text(textc)
        return textc

    #分词，返回分词结果和文本
    def readdoc3(self,path):
        doc1 = docx.Document(path)
        textc = ''
        for i in range(len(doc1.paragraphs)):
            text1 = doc1.paragraphs[i].text
            textc += text1
        textc = self.deal_text(textc)
        textwb = jieba.cut(textc, cut_all=False)
        str_out = ' '.join(textwb).replace('，', '').replace('。', '').replace('？', '').replace('！', '') \
            .replace('“', '').replace('”', '').replace('：', '').replace('…', '').replace('（', '').replace('）', '') \
            .replace('—', '').replace('《', '').replace('》', '').replace('、', '').replace('‘', '') \
            .replace('’', '')
        list = str_out.split(' ')
        while '' in list:
            list.remove('')
        return list, textc

    #得到分词后的文本list
    def get_doclist(self,path, file):
        filename = os.path.splitext(file)[0]
        doc1 = docx.Document(os.path.join(path, file))

        doc_fenci = self.readdoc(doc1)
        return doc_fenci, filename

    #以dict的形式返回每个目录下的文件数{文件名：文件数}
    def get_filecount(self,paths):
        count = {}
        for path in paths:
            count[path] = len(os.listdir(path))
        return count

    #读取csv文件
    def readdata(self, path):
        data = pd.read_csv(path, encoding='gbk')
        return data




